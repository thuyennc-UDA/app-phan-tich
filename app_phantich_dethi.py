import streamlit as st
import pandas as pd
import numpy as np
from io import BytesIO
from datetime import datetime

# ====== Các hàm tính toán ======
def cronbach_alpha(df_items):
    if df_items.shape[1] < 2:
        return np.nan
    df_corr = df_items.corr()
    N = df_corr.shape[0]
    mean_r = df_corr.values[np.triu_indices(N, 1)].mean()
    alpha = (N * mean_r) / (1 + (N - 1) * mean_r)
    return alpha

def classify_difficulty(p):
    if p >= 0.8: return "Dễ"
    elif p >= 0.6: return "Trung bình"
    elif p >= 0.4: return "Tương đối khó"
    elif p >= 0.2: return "Khó"
    else: return "Rất khó"

def classify_discrimination(d):
    if d >= 0.7: return "Rất tốt"
    elif d >= 0.4: return "Tốt"
    elif d >= 0.2: return "Đủ phân biệt"
    elif d >= 0: return "Yếu"
    else: return "Rất yếu (nên loại bỏ)"

def classify_reliability(r):
    if r >= 0.85:
        return "Đề thi có độ tin cậy rất cao"
    elif r >= 0.80:
        return "Đề thi có độ tin cậy cao"
    elif r >= 0.70:
        return "Đề thi đủ độ tin cậy"
    elif r >= 0.60:
        return "Đề thi có độ tin cậy yếu"
    else:
        return "Đề thi không đủ độ tin cậy"

# ====== Giao diện ======
st.set_page_config(page_title="Phân tích điểm thi", layout="wide")
st.title("📊 Ứng dụng phân tích điểm thi theo mã đề")

uploaded_file = st.file_uploader(
    "Tải file CSV/XLSX (có cột 'Tên học phần' và 'Mã đề')", 
    type=["csv", "xlsx"]
)

if uploaded_file:
    # Đọc file
    if uploaded_file.name.endswith("csv"):
        df = pd.read_csv(uploaded_file)
    else:
        df = pd.read_excel(uploaded_file)

    st.subheader("📑 Dữ liệu gốc")
    st.dataframe(df.head(10))

    hocphan = st.selectbox("Chọn học phần", df["Tên học phần"].unique())
    df_hp = df[df["Tên học phần"] == hocphan]

    all_item_cols = [c for c in df_hp.columns if c.lower().startswith("cau")]
    so_cot = len(all_item_cols)

    so_cau = st.number_input(
        f"Số câu hỏi của học phần {hocphan}", 
        min_value=1, max_value=so_cot, value=so_cot, step=1
    )
    item_cols = all_item_cols[:so_cau]

    st.write("Nhập điểm tối đa cho từng câu:")
    max_scores = []
    cols = st.columns(min(so_cau, 5))
    for i, col_name in enumerate(item_cols):
        col = cols[i % len(cols)]
        max_val = col.number_input(f"{col_name}", min_value=0.5, max_value=20.0, value=4.0, step=0.5)
        max_scores.append(max_val)
    max_scores = pd.Series(max_scores, index=item_cols)

    # ====== Nút chạy phân tích ======
    if st.button("🚀 Chạy phân tích dữ liệu"):
        results_dict = {}

        for made in df_hp["Made"].unique():
            df_md = df_hp[df_hp["Made"] == made]
            df_items = df_md[item_cols]

            # Độ khó
            df_diff = pd.DataFrame({
                "Câu": item_cols,
                "Điểm tối đa": max_scores.values,
                "Điểm trung bình": df_items.mean().values,
                "Độ khó": (df_items.mean() / max_scores).values,
                "Mức độ": [classify_difficulty(p) for p in (df_items.mean() / max_scores).values]
            })

            # Độ phân biệt
            n_students = len(df_items)
            group_size = max(1, int(round(0.27 * n_students)))
            df_sorted = df_md.sort_values("TongDiem", ascending=False)
            top_group = df_sorted.head(group_size)
            bottom_group = df_sorted.tail(group_size)

            top_means = top_group[item_cols].mean().values
            bottom_means = bottom_group[item_cols].mean().values

            df_disc = pd.DataFrame({
                "Câu": item_cols,
                "Điểm tối đa": max_scores.values,
                "Số thí sinh mỗi nhóm": [group_size] * len(item_cols),
                "Điểm TB nhóm trên": top_means,
                "Điểm TB nhóm dưới": bottom_means,
                "Độ phân biệt": (top_means - bottom_means) / max_scores.values,
                "Mức độ": [classify_discrimination(d) for d in (top_means - bottom_means) / max_scores.values]
            })

            # Độ tin cậy
            items = df_items.values
            k = items.shape[1]
            item_vars = df_items.var(ddof=1)
            total_var = df_items.sum(axis=1).var(ddof=1)
            alpha = (k / (k - 1)) * (1 - item_vars.sum() / total_var)
            alpha = round(alpha, 2)

            df_reliability = pd.DataFrame({
                "Câu hỏi": list(item_cols) + ["Tổng điểm"],
                "Phương sai của câu": list(item_vars.values) + [total_var]
            })
            df_reliability.loc[len(df_reliability)] = ["Tổng phương sai các câu hỏi", item_vars.sum()]
            df_reliability.loc[len(df_reliability)] = ["Cronbach's Alpha", alpha]
            df_reliability.loc[len(df_reliability)] = ["Đánh giá", classify_reliability(alpha)]

            results_dict[made] = {
                "Độ khó": df_diff,
                "Phân biệt": df_disc,
                "Tin cậy": df_reliability
            }

            # Hiển thị trên giao diện
            st.markdown(f"## 📘 MÃ ĐỀ {made}")
            st.markdown("### 1️⃣ Độ khó")
            st.dataframe(df_diff)
            st.markdown("### 2️⃣ Độ phân biệt")
            st.dataframe(df_disc)
            st.markdown("### 3️⃣ Độ tin cậy (Cronbach's Alpha)")
            st.dataframe(df_reliability)

        st.session_state["results_dict"] = results_dict
        st.success("✅ Đã hoàn thành phân tích dữ liệu!")

# ====== Xuất Excel ======
st.markdown("---")
st.subheader("📤 Xuất kết quả ra Excel")

if st.button("💾 Xuất file Excel"):
    results_dict = st.session_state.get("results_dict", None)
    if not results_dict:
        st.warning("⚠️ Vui lòng chạy phân tích trước khi xuất file Excel!")
    else:
        output = BytesIO()
        with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
            for made, result in results_dict.items():
                for name, df_part in result.items():
                    df_part.to_excel(writer, index=False, sheet_name=f"{made}_{name}")
        output.seek(0)
        filename = f"PhanTich_DeThi_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
        st.download_button(
            label="📥 Tải file Excel",
            data=output.getvalue(),
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        st.success("✅ Đã tạo file Excel thành công!")

# ====== Xuất Word ======
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.markdown("---")
st.subheader("📝 Xuất báo cáo ra Word")

de_xuat_text = st.text_area(
    "✏️ Nhập nội dung phần III. Đề xuất, cải tiến:",
    "VÍ DỤ:  Đề xuất điều chỉnh độ khó của một số câu hỏi để cân bằng mức độ đánh giá năng lực sinh viên. "
    "Bổ sung thêm các câu hỏi kiểm tra mức vận dụng và phân tích để cải thiện độ phân biệt của đề thi."
)

if st.button("📄 Xuất báo cáo Word"):
    results_dict = st.session_state.get("results_dict", None)
    if not results_dict:
        st.warning("⚠️ Vui lòng chạy phân tích trước khi xuất báo cáo Word!")
    else:
        doc = Document()

        # ===== Cấu hình trang A4 =====
        section = doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        # ===== Font mặc định =====
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10)

        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.shared import OxmlElement, qn
        import os

        # --- Tạo document ---
        doc = Document()

        # --- CHÈN LOGO + TIÊU ĐỀ SONG SONG ---
        # Tạo bảng 1 hàng 2 cột (để dễ căn logo & chữ)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        # Cột trái: logo
        cell_logo = table.cell(0, 0)
        logo_path = os.path.join(os.path.dirname(__file__), "logo.png")
        paragraph_logo = cell_logo.paragraphs[0]
        run_logo = paragraph_logo.add_run()
        run_logo.add_picture(logo_path, width=Inches(1.3))  # chỉnh kích thước phù hợp
        paragraph_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Cột phải: tiêu đề
        cell_title = table.cell(0, 1)
        p_title = cell_title.paragraphs[0]
        run_title = p_title.add_run("CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
        run_title.bold = True
        run_title.font.size = Pt(12)
        run_title.font.name = "Times New Roman"

        run_sub = p_title.add_run("Độc lập - Tự do - Hạnh phúc\n")
        run_sub.italic = True
        run_sub.font.size = Pt(12)
        run_sub.font.name = "Times New Roman"

        run_report = p_title.add_run("\n")
        run_report.bold = True
        run_report.font.size = Pt(13)
        run_report.font.name = "Times New Roman"
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        
        # ===== TIÊU ĐỀ =====
        title = doc.add_heading("BÁO CÁO PHÂN TÍCH ĐỀ THI", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(" ")

        from docx.shared import Pt

        # Đặt kiểu mặc định cho toàn tài liệu
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 1.2
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)

        # ===== Cấu hình trang A4 =====
        section = doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        # ===== Font mặc định =====
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10)

        # ===== PHẦN I: THÔNG TIN CHUNG =====
        doc.add_heading("I. Thông tin chung", level=2)
        doc.add_paragraph(f"1. Tên học phần: {hocphan}")
        doc.add_paragraph(f"2. Số lượng sinh viên dự thi: {len(df_hp)}")
        doc.add_paragraph(f"3. Số lượng đề thi: {df_hp['Made'].nunique()}")

        # 👉 Tính số lượng sinh viên làm mỗi mã đề
        doc.add_paragraph("4. Số lượng sinh viên làm từng mã đề:")
        table_made = doc.add_table(rows=1, cols=2)
        hdr_cells = table_made.rows[0].cells
        hdr_cells[0].text = "Mã đề"
        hdr_cells[1].text = "Số lượng sinh viên"

        made_counts = df_hp["Made"].value_counts()
        for made, count in made_counts.items():
            row_cells = table_made.add_row().cells
            row_cells[0].text = str(made)
            row_cells[1].text = str(count)

        # Tiếp theo là phần số câu
        doc.add_paragraph(f"5. Số câu trong đề thi: {so_cau}")
        doc.add_paragraph("   Điểm tối đa của từng câu hỏi:")
        table_info = doc.add_table(rows=1, cols=2)
        hdr_cells = table_info.rows[0].cells
        hdr_cells[0].text = "Câu hỏi"
        hdr_cells[1].text = "Điểm tối đa"
        for cau, diem in max_scores.items():
            row_cells = table_info.add_row().cells
            row_cells[0].text = cau
            row_cells[1].text = str(diem)
        doc.add_paragraph(" ")


        # ===== PHẦN II: PHÂN TÍCH =====
        doc.add_heading("II. Kết quả phân tích", level=2)
        for made, result in results_dict.items():
            doc.add_heading(f"MÃ ĐỀ {made}", level=3)
            
            # --- Độ khó ---
            doc.add_heading("1️⃣ Bảng phân tích độ khó", level=4)
            df_diff = result["Độ khó"].round(2)
            table_diff = doc.add_table(rows=1, cols=len(df_diff.columns))
            for i, col in enumerate(df_diff.columns):
                table_diff.rows[0].cells[i].text = col
            for _, row in df_diff.iterrows():
                cells = table_diff.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)
            doc.add_paragraph(" ")

            # --- Độ phân biệt ---
            doc.add_heading("2️⃣ Bảng phân tích độ phân biệt", level=4)
            df_disc = result["Phân biệt"].round(2)
            table_disc = doc.add_table(rows=1, cols=len(df_disc.columns))
            for i, col in enumerate(df_disc.columns):
                table_disc.rows[0].cells[i].text = col
            for _, row in df_disc.iterrows():
                cells = table_disc.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)
            doc.add_paragraph(" ")

            # --- Độ tin cậy ---
            doc.add_heading("3️⃣ Bảng phân tích độ tin cậy", level=4)
            df_rel = result["Tin cậy"].round(2)
            table_rel = doc.add_table(rows=1, cols=len(df_rel.columns))
            for i, col in enumerate(df_rel.columns):
                table_rel.rows[0].cells[i].text = col
            for _, row in df_rel.iterrows():
                cells = table_rel.add_row().cells
                for i, val in enumerate(row):
                    cells[i].text = str(val)
            doc.add_paragraph(" ")

        # ===== PHẦN III: ĐỀ XUẤT, CẢI TIẾN =====
        doc.add_heading("III. Đề xuất, cải tiến", level=2)
        doc.add_paragraph(de_xuat_text)

        # ===== PHỤ LỤC: BẢNG PHÂN LOẠI ====
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
            
        # ===== CHÈN HÌNH Ở DƯỚI CÙNG =====

        try:
            doc.add_heading("PHỤ LỤC: BẢNG PHÂN LOẠI ĐỘ KHÓ, ĐỘ PHÂN BIỆT, ĐỘ TIN CẬY", level=4)
            doc.add_paragraph()  # Tạo khoảng trống trước hình
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()

            # Chèn hình (tự động co cho vừa chiều ngang A4)
            import os
            img_path = os.path.join(os.path.dirname(__file__), "bangphanloai.png")
            run.add_picture(img_path, width=Inches(6.0))

            doc.add_paragraph("NGƯỜI BÁO CÁO", style='Normal').alignment = WD_ALIGN_PARAGRAPH.RIGHT

        except Exception as e:
            st.warning(f"⚠️ Không thể chèn hình bangphanloai.png: {e}")


        # ===== LƯU FILE =====
        output = BytesIO()
        doc.save(output)
        output.seek(0)
        filename = f"BaoCao_PT_Dethi_{hocphan}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        st.download_button(
            label="📥 Tải xuống báo cáo Word",
            data=output,
            file_name=filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        st.success("✅ Đã tạo file báo cáo Word thành công!")



